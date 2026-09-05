"""
ResuMap — unified FastAPI backend.

Single entry point that exposes:
  - Public anonymous scan:    POST /scan
  - Auth:                     POST /api/auth/signup, /api/auth/login,
                              POST /api/auth/forgot-password,
                              POST /api/auth/reset-password
  - User:                     GET  /api/users/me
  - Resumes:                  GET  /api/resumes, POST /api/resumes/upload,
                              DELETE /api/resumes/{id}
  - Analyses:                 POST /api/analyses, GET /api/analyses,
                              GET /api/analyses/{id},
                              DELETE /api/analyses/{id}
  - Health:                   GET  /

Run locally:
    uvicorn main:app --reload --host 0.0.0.0 --port 8000
"""
import io
import json
import os
import uuid
from contextlib import asynccontextmanager
from datetime import datetime, timezone
from typing import List

import pdfplumber
from dotenv import load_dotenv
from fastapi import (
    APIRouter, Depends, FastAPI, File, Form, HTTPException, Request, UploadFile,
    status,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.errors import RateLimitExceeded
from sqlalchemy.orm import Session

from ai import analyze_resume
from auth.dependencies import get_current_user
from auth.utils import (
    _hash_reset_token, create_access_token, generate_reset_token,
    get_password_hash, reset_token_matches, verify_password,
)
from database import get_db, init_db
from models import (
    PasswordResetToken, Resume, ResumeAnalysis, Subscription, User,
)
from resume_parser import extract_skills
from schemas import (
    CompareAnalysesRequest, CreateAnalysisRequest, ForgotPasswordRequest, Recommendation,
    ResetPasswordRequest, ResumeUploadResponse, Token, UserCreate, UserOut,
)

load_dotenv()


# ---------------------------------------------------------------------------
# App + lifespan
# ---------------------------------------------------------------------------
@asynccontextmanager
async def _lifespan(_app: FastAPI):  # noqa: ARG001
    # Create tables if they don't exist (safe to call repeatedly).
    try:
        init_db()
    except Exception as exc:  # pragma: no cover
        # Log but don't crash — endpoints that don't need the DB still work.
        print(f"[startup] init_db failed: {exc}")
    yield


app = FastAPI(
    title="ResuMap — AI Resume Screener",
    version="2.0.0",
    description="Unified backend: auth, resume storage, AI scan.",
    lifespan=_lifespan,
)

# Allowed browser origins — never "*". Override via CORS_ORIGINS (comma-separated)
# when a custom frontend domain is added. Dev origins + the production Vercel app.
_DEFAULT_CORS = (
    "http://localhost:5173,http://127.0.0.1:5173,https://resume-screener-one.vercel.app"
)
CORS_ORIGINS = [
    o.strip() for o in os.getenv("CORS_ORIGINS", _DEFAULT_CORS).split(",") if o.strip()
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Rate limiting (slowapi). In-memory store is fine for a single free instance;
# the point is to stop anonymous Groq abuse and auth brute-force, not global
# fairness. Real client IP comes from Render's X-Forwarded-For proxy header.
# ---------------------------------------------------------------------------
def _client_ip(request: Request) -> str:
    forwarded = request.headers.get("x-forwarded-for")
    if forwarded:
        return forwarded.split(",")[0].strip()
    return request.client.host if request.client else "unknown"


limiter = Limiter(key_func=_client_ip)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

ALLOWED_RESUME_EXT = {".pdf", ".docx"}
MAX_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB


# ---------------------------------------------------------------------------
# Public anonymous scan (no auth) — used by the legacy /scan and the
# "Try without signup" flow.
# ---------------------------------------------------------------------------
public_router = APIRouter()


@public_router.post("/scan")
@limiter.limit("10/minute")
async def scan_resume(
    request: Request,
    file: UploadFile = File(...),
    job_description: str = Form(...),
):
    """Anonymous single-shot scan. Returns AI analysis only — does not save."""
    if not file.filename or not any(
        file.filename.lower().endswith(ext) for ext in ALLOWED_RESUME_EXT
    ):
        raise HTTPException(
            status_code=400, detail="Only PDF and DOCX files are allowed"
        )

    try:
        content = await file.read()
        if len(content) > MAX_UPLOAD_BYTES:
            raise HTTPException(status_code=413, detail="File too large (max 10MB)")

        if file.filename.lower().endswith(".pdf"):
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                resume_text = "".join(
                    (page.extract_text() or "") for page in pdf.pages
                )
        else:  # .docx
            from docx import Document
            doc = Document(io.BytesIO(content))
            resume_text = "\n".join(p.text for p in doc.paragraphs)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(
            status_code=500, detail=f"File extraction failed: {exc}"
        )

    if not resume_text.strip():
        raise HTTPException(
            status_code=400,
            detail="Could not extract text from the file. It may be image-only.",
        )

    try:
        result = analyze_resume(resume_text, job_description)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {exc}")

    return result


# ---------------------------------------------------------------------------
# Auth
# ---------------------------------------------------------------------------
auth_router = APIRouter(prefix="/api/auth", tags=["auth"])


@auth_router.post("/signup", response_model=Token, status_code=status.HTTP_201_CREATED)
@limiter.limit("20/minute")
def signup(request: Request, payload: UserCreate, db: Session = Depends(get_db)):
    if len(payload.password) < 6:
        raise HTTPException(
            status_code=400,
            detail="Password must be at least 6 characters",
        )
    existing = db.query(User).filter(User.email == payload.email).first()
    if existing:
        raise HTTPException(status_code=409, detail="Email already exists")

    user = User(
        email=payload.email,
        password_hash=get_password_hash(payload.password),
        full_name=payload.full_name,
    )
    db.add(user)
    db.flush()  # populate user.id

    # Free-tier subscription by default.
    db.add(Subscription(user_id=user.id, plan_type="free"))

    db.commit()
    db.refresh(user)

    token = create_access_token({"sub": str(user.id)})
    return Token(access_token=token, user_id=str(user.id))


@auth_router.post("/login", response_model=Token)
@limiter.limit("20/minute")
def login(request: Request, payload: UserCreate, db: Session = Depends(get_db)):
    user = db.query(User).filter(User.email == payload.email).first()
    if not user or not verify_password(payload.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    token = create_access_token({"sub": str(user.id)})
    return Token(access_token=token, user_id=str(user.id))


# Generic response for forgot-password — same message whether or not the
# email exists, so an attacker can't enumerate which addresses are signed up.
_RESET_RESPONSE = {"message": "If that email exists, a reset link has been sent."}


@auth_router.post("/forgot-password")
@limiter.limit("5/minute")
def forgot_password(
    request: Request,
    payload: ForgotPasswordRequest,
    db: Session = Depends(get_db),
):
    """Issue a one-time password reset token.

    Always returns the same generic response, regardless of whether the email
    exists, to prevent account enumeration. When the email is known, a fresh
    token is generated, all previous tokens for that user are invalidated, and
    the plaintext token is logged (and would normally be emailed — see comment
    below).
    """
    user = db.query(User).filter(User.email == payload.email).first()
    if user:
        # Invalidate any existing unused tokens for this user so only the most
        # recent link works.
        db.query(PasswordResetToken).filter(
            PasswordResetToken.user_id == user.id
        ).delete(synchronize_session=False)

        raw_token, token_hash, expires_at = generate_reset_token()
        db.add(PasswordResetToken(
            user_id=user.id,
            token_hash=token_hash,
            expires_at=expires_at,
        ))
        db.commit()

        # In production this would be an SES/SendGrid/SMTP call. To keep the
        # app self-contained and testable without an email provider, we log
        # the link instead. Swap this for an actual email send when SMTP is
        # configured.
        print(
            f"[reset] Password reset link for {user.email}: "
            f"/reset-password?token={raw_token}"
        )

    return _RESET_RESPONSE


@auth_router.post("/reset-password")
def reset_password(
    payload: ResetPasswordRequest,
    db: Session = Depends(get_db),
):
    """Consume a reset token and set a new password.

    Validates the token (exists, not expired, hash matches), updates the user's
    password, and deletes the token so it can't be reused. Any other live
    tokens for that user are also cleared so the old link stops working too.
    """
    if len(payload.new_password) < 6:
        raise HTTPException(
            status_code=400,
            detail="Password must be at least 6 characters",
        )

    # Compute the hash of the incoming token the same way generate_reset_token
    # stored it (sha256 hex), then look it up.
    from auth.utils import _hash_reset_token
    token_hash = _hash_reset_token(payload.token)

    record = (
        db.query(PasswordResetToken)
        .filter(PasswordResetToken.token_hash == token_hash)
        .first()
    )
    if not record:
        raise HTTPException(
            status_code=400,
            detail="Invalid or expired reset link",
        )

    # Constant-time check (matches the unique-index lookup above, but kept so
    # timing-side-channel review is honest if the index ever changes).
    if not reset_token_matches(payload.token, record.token_hash):
        raise HTTPException(
            status_code=400,
            detail="Invalid or expired reset link",
        )

    # SQLite strips tz info; normalize so the comparison is correct on both
    # backends.
    expires_at = record.expires_at
    if expires_at.tzinfo is None:
        expires_at = expires_at.replace(tzinfo=timezone.utc)
    if expires_at < datetime.now(timezone.utc):
        db.delete(record)
        db.commit()
        raise HTTPException(
            status_code=400,
            detail="Invalid or expired reset link",
        )

    user = db.query(User).filter(User.id == record.user_id).first()
    if not user:
        # FK cascade should prevent this, but guard anyway.
        db.delete(record)
        db.commit()
        raise HTTPException(
            status_code=400,
            detail="Invalid or expired reset link",
        )

    user.password_hash = get_password_hash(payload.new_password)
    # Invalidate ALL reset tokens for this user — the one we just used and any
    # stragglers from earlier forgot-password calls.
    db.query(PasswordResetToken).filter(
        PasswordResetToken.user_id == user.id
    ).delete(synchronize_session=False)
    db.commit()

    return {"message": "Password updated. You can now sign in."}


# ---------------------------------------------------------------------------
# User
# ---------------------------------------------------------------------------
user_router = APIRouter(prefix="/api/users", tags=["user"])


@user_router.get("/me", response_model=UserOut)
def get_me(current: User = Depends(get_current_user)):
    return current


# ---------------------------------------------------------------------------
# Resumes
# ---------------------------------------------------------------------------
resume_router = APIRouter(prefix="/api/resumes", tags=["resumes"])


def _store_file(user_id: str, original_filename: str, content: bytes) -> str:
    """Persist an upload and return its on-disk path."""
    ext = os.path.splitext(original_filename)[1].lower() or ".pdf"
    safe_name = f"{user_id}_{uuid.uuid4().hex}{ext}"
    path = os.path.join(UPLOAD_DIR, safe_name)
    with open(path, "wb") as fh:
        fh.write(content)
    return path


def _extract_text(filename: str, content: bytes) -> str:
    if filename.lower().endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            return "".join((page.extract_text() or "") for page in pdf.pages)
    if filename.lower().endswith(".docx"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)
    raise HTTPException(status_code=400, detail="Unsupported file type")


@resume_router.get("", response_model=List[dict])
def list_resumes(
    current: User = Depends(get_current_user), db: Session = Depends(get_db)
):
    rows = (
        db.query(Resume)
        .filter(Resume.user_id == current.id)
        .order_by(Resume.created_at.desc())
        .all()
    )
    return [
        {
            "id": str(r.id),
            "file_name": r.file_name,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in rows
    ]


@resume_router.post(
    "/upload", response_model=ResumeUploadResponse, status_code=status.HTTP_201_CREATED
)
async def upload_resume(
    file: UploadFile = File(...),
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    if not file.filename or not any(
        file.filename.lower().endswith(ext) for ext in ALLOWED_RESUME_EXT
    ):
        raise HTTPException(
            status_code=400, detail="Only PDF and DOCX files are allowed"
        )

    content = await file.read()
    if len(content) > MAX_UPLOAD_BYTES:
        raise HTTPException(status_code=413, detail="File too large (max 10MB)")

    try:
        text = _extract_text(file.filename, content)
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Extraction failed: {exc}")

    skills = extract_skills(text)
    path = _store_file(str(current.id), file.filename, content)

    # Mark previous resumes as not current.
    db.query(Resume).filter(
        Resume.user_id == current.id, Resume.is_current.is_(True)
    ).update({Resume.is_current: False})

    resume = Resume(
        user_id=current.id,
        file_name=file.filename,
        file_path=path,
        raw_text=text,
        is_current=True,
    )
    db.add(resume)
    db.commit()
    db.refresh(resume)

    return ResumeUploadResponse(
        resume_id=str(resume.id),
        skills=skills,
        word_count=len(text.split()),
    )


@resume_router.delete("/{resume_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_resume(
    resume_id: str,
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete a resume owned by the current user. Cascades to its analyses."""
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current.id)
        .first()
    )
    if not resume:
        # Same code for missing and not-yours — don't leak ownership info.
        raise HTTPException(status_code=404, detail="Resume not found")

    was_current = bool(resume.is_current)

    # Best-effort filesystem cleanup. We don't fail the request if the file is
    # already gone (ephemeral Render disk loses files on restart anyway).
    file_path = resume.file_path
    db.delete(resume)
    db.flush()

    if was_current:
        # Promote the next-most-recent remaining resume so "is_current"
        # invariant holds (at most one current resume per user).
        next_resume = (
            db.query(Resume)
            .filter(Resume.user_id == current.id)
            .order_by(Resume.created_at.desc())
            .first()
        )
        if next_resume is not None:
            next_resume.is_current = True

    db.commit()

    if file_path:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except OSError as exc:  # pragma: no cover
            print(f"[delete_resume] could not remove {file_path}: {exc}")


@resume_router.get("/{resume_id}/download")
def download_resume(
    resume_id: str,
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Serve the originally uploaded file back to its owner.

    Used by the Dashboard preview/download affordance. Render's ephemeral disk
    means a file may be gone between upload and download — in that case we
    return 410 Gone (the metadata is still there, the bytes are not). For
    privacy, missing-and-not-yours both return 404 (no existence leak).
    """
    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current.id)
        .first()
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    file_path = resume.file_path
    if not file_path or not os.path.exists(file_path):
        # Common on Render: disk is ephemeral, files vanish on restart.
        raise HTTPException(
            status_code=410,
            detail=(
                "Original file is no longer available. Re-upload your resume "
                "to restore download/preview."
            ),
        )

    # Use the original extension (PDF/DOCX) for the served name and pick the
    # right MIME type so browsers open PDFs inline instead of forcing a save.
    ext = os.path.splitext(file_path)[1].lower()
    media_type = (
        "application/pdf"
        if ext == ".pdf"
        else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        if ext == ".docx"
        else "application/octet-stream"
    )
    download_name = resume.file_name or f"resume{ext}"
    return FileResponse(
        path=file_path,
        media_type=media_type,
        filename=download_name,
    )


# ---------------------------------------------------------------------------
# Analyses
# ---------------------------------------------------------------------------
analysis_router = APIRouter(prefix="/api/analyses", tags=["analyses"])


def _to_recommendations(suggestions: List[dict]) -> List[Recommendation]:
    return [
        Recommendation(
            type=s.get("type", "gap"),
            priority=s.get("priority", "medium"),
            text=s.get("text", ""),
            action=s.get("action", "Update resume")
        )
        for s in suggestions
    ]


@analysis_router.post("", status_code=status.HTTP_201_CREATED)
def create_analysis(
    payload: CreateAnalysisRequest,
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    resume_id = payload.resume_id
    job_title = payload.job_title
    job_description = payload.job_description

    resume = (
        db.query(Resume)
        .filter(Resume.id == resume_id, Resume.user_id == current.id)
        .first()
    )
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    try:
        result = analyze_resume(resume.raw_text or "", job_description)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {exc}")

    score = float(result.get("overall_score", 0) or 0)
    missing = result.get("missing_keywords", []) or []
    suggestions = result.get("top_suggestions", []) or []
    recommendations = _to_recommendations(suggestions)
    verdict = result.get("verdict", "")

    analysis = ResumeAnalysis(
        resume_id=resume.id,
        job_title=job_title,
        job_description=job_description,
        match_score=score,
        missing_skills=missing,
        recommendations=[r.model_dump() for r in recommendations],
        verdict=verdict,
    )
    db.add(analysis)
    db.commit()
    db.refresh(analysis)

    return {
        "analysis_id": str(analysis.id),
        "resume_id": str(resume.id),
        "job_title": job_title,
        "match_score": score,
        "missing_skills": missing,
        "recommendations": [r.model_dump() for r in recommendations],
        "verdict": verdict,
        "created_at": analysis.generated_at.isoformat() if analysis.generated_at else None,
    }


@analysis_router.get("")
def list_analyses(
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """List all analyses owned by the current user (most recent first)."""
    rows = (
        db.query(ResumeAnalysis)
        .join(Resume, ResumeAnalysis.resume_id == Resume.id)
        .filter(Resume.user_id == current.id)
        .order_by(ResumeAnalysis.generated_at.desc())
        .all()
    )

    def _maybe_load(value):
        if isinstance(value, str):
            try:
                return json.loads(value)
            except Exception:
                return value
        return value

    return [
        {
            "id": str(a.id),
            "resume_id": str(a.resume_id),
            "job_title": a.job_title,
            "match_score": float(a.match_score or 0),
            "missing_skills": _maybe_load(a.missing_skills) or [],
            "recommendations": _maybe_load(a.recommendations) or [],
            "verdict": a.verdict,
            "generated_at": a.generated_at.isoformat() if a.generated_at else None,
        }
        for a in rows
    ]


@analysis_router.get("/{analysis_id}")
def get_analysis(
    analysis_id: str,
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    analysis = (
        db.query(ResumeAnalysis)
        .join(Resume, ResumeAnalysis.resume_id == Resume.id)
        .filter(
            ResumeAnalysis.id == analysis_id, Resume.user_id == current.id
        )
        .first()
    )
    if not analysis:
        raise HTTPException(status_code=404, detail="Analysis not found")

    def _maybe_load(value):
        if isinstance(value, str):
            try:
                return json.loads(value)
            except Exception:
                return value
        return value

    return {
        "id": str(analysis.id),
        "resume_id": str(analysis.resume_id),
        "job_title": analysis.job_title,
        "job_description": analysis.job_description,
        "match_score": float(analysis.match_score or 0),
        "missing_skills": _maybe_load(analysis.missing_skills) or [],
        "recommendations": _maybe_load(analysis.recommendations) or [],
        "verdict": analysis.verdict,
        "generated_at": analysis.generated_at.isoformat() if analysis.generated_at else None,
    }


@analysis_router.post("/compare")
def compare_analyses(
    payload: CompareAnalysesRequest,
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Return a side-by-side comparison of two or more analyses owned by the
    current user. Builds a per-role view the Dashboard uses to show "which
    skills are valued where" — each analysis is its own role, scored
    independently against the same resume.

    Body: { "analysis_ids": ["<id>", "<id>", ...] }   (min 2, max 6)

    Response:
      {
        "analyses": [
          {
            "id": "...",
            "job_title": "...",
            "match_score": 72.5,
            "missing_skills": ["..."],
            "verdict": "...",
            "generated_at": "...",
            "resume_id": "...",
            "resume_file_name": "..."
          },
          ...
        ],
        "skill_matrix": [
          {
            "skill": "kubernetes",
            "missing_in": ["<analysis_id>", ...]   // which roles flagged it as missing
          },
          ...
        ]
      }

    Ownership: every id must belong to the current user; if any do not, 404
    (same as the rest of the API — don't leak existence).
    """
    ids = payload.analysis_ids
    if len(ids) < 2:
        raise HTTPException(
            status_code=400,
            detail="analysis_ids must be a list of at least 2 ids",
        )
    if len(ids) > 6:
        raise HTTPException(
            status_code=400,
            detail="compare at most 6 analyses at once",
        )
    # Dedupe while preserving order, so duplicate ids don't break the matrix.
    seen = set()
    unique_ids = []
    for aid in ids:
        if aid not in seen:
            seen.add(aid)
            unique_ids.append(aid)
    if len(unique_ids) < 2:
        raise HTTPException(
            status_code=400,
            detail="need at least 2 distinct analyses to compare",
        )

    def _maybe_load(value):
        if isinstance(value, str):
            try:
                return json.loads(value)
            except Exception:
                return value
        return value

    rows = (
        db.query(ResumeAnalysis, Resume.file_name)
        .join(Resume, ResumeAnalysis.resume_id == Resume.id)
        .filter(
            ResumeAnalysis.id.in_(unique_ids),
            Resume.user_id == current.id,
        )
        .all()
    )
    if len(rows) != len(unique_ids):
        # Some ids were missing or not owned by this user.
        raise HTTPException(
            status_code=404,
            detail="One or more analyses not found",
        )

    # Build the response in the order the client requested (so the columns in
    # the matrix line up with the cards the user just clicked).
    by_id = {str(a.id): (a, fname) for a, fname in rows}
    analyses_out = []
    missing_by_id = {}
    for aid in unique_ids:
        a, fname = by_id[aid]
        missing = _maybe_load(a.missing_skills) or []
        if not isinstance(missing, list):
            missing = []
        # Normalize: lowercase, strip, drop empties — so "Kubernetes" and
        # "kubernetes" count as the same skill in the matrix.
        normalized = []
        for s in missing:
            if not isinstance(s, str):
                continue
            n = s.strip().lower()
            if n and n not in normalized:
                normalized.append(n)
        missing_by_id[aid] = normalized
        analyses_out.append({
            "id": aid,
            "resume_id": str(a.resume_id),
            "resume_file_name": fname,
            "job_title": a.job_title or "Untitled role",
            "match_score": float(a.match_score or 0),
            "missing_skills": missing,
            "verdict": a.verdict or "",
            "generated_at": a.generated_at.isoformat() if a.generated_at else None,
        })

    # Build the skill matrix: every distinct missing skill across the
    # selected analyses, sorted by how many roles flagged it (desc) so the
    # "most universally required" skills surface first.
    counts: dict[str, int] = {}
    for skills in missing_by_id.values():
        for s in skills:
            counts[s] = counts.get(s, 0) + 1
    # Stable secondary sort by skill name so ties don't reshuffle on every
    # request.
    sorted_skills = sorted(counts.items(), key=lambda kv: (-kv[1], kv[0]))
    skill_matrix = [
        {
            "skill": skill,
            "missing_in": [
                aid for aid in unique_ids if skill in missing_by_id[aid]
            ],
        }
        for skill, _ in sorted_skills
    ]

    return {"analyses": analyses_out, "skill_matrix": skill_matrix}


@analysis_router.delete("/{analysis_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_analysis(
    analysis_id: str,
    current: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete an analysis owned by the current user. Does NOT delete the resume."""
    analysis = (
        db.query(ResumeAnalysis)
        .join(Resume, ResumeAnalysis.resume_id == Resume.id)
        .filter(
            ResumeAnalysis.id == analysis_id, Resume.user_id == current.id
        )
        .first()
    )
    if not analysis:
        # Same code for missing and not-yours — don't leak ownership info.
        raise HTTPException(status_code=404, detail="Analysis not found")

    db.delete(analysis)
    db.commit()


# ---------------------------------------------------------------------------
# Health
# ---------------------------------------------------------------------------
@app.get("/")
def root():
    return {
        "service": "ResuMap API",
        "version": app.version,
        "status": "ok",
        "endpoints": [
            "POST /scan",
            "POST /api/auth/signup",
            "POST /api/auth/login",
            "GET  /api/users/me",
            "GET  /api/resumes",
            "POST /api/resumes/upload",
            "GET  /api/resumes/{id}/download",
            "DELETE /api/resumes/{id}",
            "POST /api/analyses",
            "GET  /api/analyses",
            "POST /api/analyses/compare",
            "GET  /api/analyses/{id}",
            "DELETE /api/analyses/{id}",
        ],
    }


app.include_router(public_router)
app.include_router(auth_router)
app.include_router(user_router)
app.include_router(resume_router)
app.include_router(analysis_router)


if __name__ == "__main__":
    import uvicorn
    uvicorn.run("main:app", host="127.0.0.1", port=5000, reload=True)
